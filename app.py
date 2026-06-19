import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
from datetime import datetime

# ─── PAGE CONFIG ──────────────────────────────────────────────────────────────
st.set_page_config(page_title="Order Studio", page_icon="📦", layout="wide")

# ─── GLOBAL STYLES — Minimalism 2026 Light Mode ──────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&family=JetBrains+Mono:wght@400;500;600&display=swap');

/* ════════════════════════════════════════════════════════════
   DESIGN TOKENS — Minimalism 2026
════════════════════════════════════════════════════════════
   Background : #FFFFFF
   Surface    : #F8FAFC
   Border     : #E2E8F0
   Text-1     : #0F172A  (primary)
   Text-2     : #475569  (secondary)
   Text-3     : #94A3B8  (muted)
   Accent     : #2563EB  (blue-600)
   Accent-bg  : #EFF6FF  (blue-50)
   Accent-bdr : #BFDBFE  (blue-200)
   Success    : #16A34A
   Warning    : #D97706
════════════════════════════════════════════════════════════ */

*, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }
html { -webkit-font-smoothing: antialiased; text-rendering: optimizeLegibility; }

.stApp {
    font-family: 'Inter', system-ui, -apple-system, sans-serif;
    background: #FFFFFF;
    min-height: 100vh;
    color: #0F172A;
}

#MainMenu, footer, header { visibility: hidden; display: none; height: 0; overflow: hidden; }
[data-testid="stAppDeployButton"] { display: none !important; }
.viewerBadge_container__r5tak,
.viewerBadge_link__qRIco { display: none !important; }
a[href*="streamlit.io/cloud"], a[href*="share.streamlit.io"] { display: none !important; }

.block-container {
    padding-top: 0 !important;
    padding-bottom: 5rem !important;
    max-width: 100% !important;
    padding-left: 2.5rem !important;
    padding-right: 2.5rem !important;
}
.stMainBlockContainer > div {
    max-width: 1100px;
    margin-left: auto;
    margin-right: auto;
}

/* ════════════════════════════════════════════════════════════
   TOPBAR
════════════════════════════════════════════════════════════ */
.os-topbar {
    display: flex;
    align-items: center;
    justify-content: space-between;
    padding: 28px 0 24px;
    margin-bottom: 8px;
    border-bottom: 1px solid #F1F5F9;
}
.os-topbar-brand { display: flex; align-items: center; gap: 12px; }
.os-topbar-logo {
    width: 36px; height: 36px;
    background: #EFF6FF;
    border: 1px solid #BFDBFE;
    border-radius: 10px;
    display: flex; align-items: center; justify-content: center;
    font-size: 17px;
    flex-shrink: 0;
}
.os-topbar-name {
    font-size: 17px;
    font-weight: 700;
    color: #0F172A;
    letter-spacing: -0.4px;
    line-height: 1.2;
}
.os-topbar-sub {
    font-size: 12px;
    color: #94A3B8;
    font-weight: 400;
    margin-top: 2px;
    letter-spacing: 0;
}
.os-topbar-status {
    display: flex; align-items: center; gap: 6px;
    font-size: 12px; font-weight: 500; color: #64748B;
    background: #F8FAFC;
    border: 1px solid #E2E8F0;
    border-radius: 100px;
    padding: 6px 14px;
    letter-spacing: 0;
}
.os-status-dot {
    width: 6px; height: 6px; border-radius: 50%;
    background: #16A34A;
    display: inline-block;
    flex-shrink: 0;
}

/* ════════════════════════════════════════════════════════════
   SECTION LABEL
════════════════════════════════════════════════════════════ */
.os-section {
    display: flex; align-items: center; gap: 10px;
    margin: 28px 0 14px;
}
.os-section-line {
    flex: 1;
    height: 1px;
    background: #F1F5F9;
}
.os-section-title {
    font-size: 11px;
    font-weight: 600;
    color: #94A3B8;
    letter-spacing: 0.8px;
    text-transform: uppercase;
    white-space: nowrap;
}
.os-section-tag {
    font-size: 10px; font-weight: 600;
    color: #2563EB;
    background: #EFF6FF;
    padding: 3px 10px; border-radius: 100px;
    text-transform: uppercase;
    letter-spacing: 0.6px;
    font-family: 'Inter', sans-serif;
    border: 1px solid #BFDBFE;
    white-space: nowrap;
}
.os-section-sep { flex: 1; }

/* ════════════════════════════════════════════════════════════
   SETTINGS SELECTS
════════════════════════════════════════════════════════════ */
.stSelectbox > div > div {
    border: 1px solid #E2E8F0 !important;
    border-radius: 10px !important;
    background: #FFFFFF !important;
    font-size: 14px !important;
    font-family: 'Inter', sans-serif !important;
    transition: border-color 0.15s ease, box-shadow 0.15s ease !important;
    box-shadow: none !important;
    color: #0F172A !important;
}
.stSelectbox > div > div:hover {
    border-color: #BFDBFE !important;
}
.stSelectbox > div > div:focus-within {
    border-color: #2563EB !important;
    box-shadow: 0 0 0 3px rgba(37,99,235,0.08) !important;
}
.stSelectbox label {
    font-size: 12px !important;
    font-weight: 500 !important;
    color: #64748B !important;
    text-transform: none !important;
    letter-spacing: 0 !important;
    margin-bottom: 6px !important;
}

/* ════════════════════════════════════════════════════════════
   FILE UPLOADER
════════════════════════════════════════════════════════════ */
[data-testid="stFileUploader"] {
    background: #FFFFFF;
    border: 1.5px dashed #D1D5DB !important;
    border-radius: 14px !important;
    padding: 0;
    transition: border-color 0.15s ease, background 0.15s ease;
    margin-top: 6px !important;
    overflow: hidden;
    height: 110px !important;
    display: flex !important;
    flex-direction: column !important;
    position: relative !important;
}
[data-testid="stFileUploader"]:hover {
    border-color: #2563EB !important;
    background: #EFF6FF;
}
[data-testid="stFileUploader"] label {
    font-size: 13px !important;
    font-weight: 500 !important;
    color: #334155 !important;
    text-transform: none !important;
    letter-spacing: 0 !important;
    text-align: left !important;
    display: block !important;
    padding: 14px 16px 0 !important;
}
[data-testid="stFileUploader"] [data-testid="stTooltipHoverTarget"] {
    display: none !important;
}
[data-testid="stFileUploaderDropzoneInstructions"] {
    font-size: 0 !important;
    line-height: 0 !important;
    height: 0 !important;
    min-height: 0 !important;
    overflow: hidden !important;
    margin: 0 !important;
    padding: 0 !important;
    visibility: hidden !important;
    position: absolute !important;
}
[data-testid="stFileUploaderDropzoneInstructions"] * { display: none !important; }

[data-testid="stFileUploaderDropzone"] {
    border: none !important;
    background: transparent !important;
    width: 100% !important;
    min-width: 0 !important;
}
[data-testid="stFileUploaderDropzone"]:has([data-testid="stFileUploaderDropzoneInstructions"]) {
    display: flex !important;
    flex-direction: column !important;
    align-items: center !important;
    justify-content: center !important;
    gap: 0 !important;
    padding: 10px 16px 16px !important;
    text-align: center !important;
    flex: 1 !important;
}
[data-testid="stFileUploaderDropzone"]:has([data-testid="stFileUploaderDropzoneInstructions"])::before {
    content: 'Kéo thả file vào đây';
    display: block !important;
    width: 100% !important;
    font-size: 12px;
    color: #94A3B8;
    font-weight: 400;
    font-family: 'Inter', sans-serif;
    text-align: center !important;
    line-height: 1;
    margin-bottom: 12px;
    order: -2;
}
[data-testid="stFileUploaderDropzone"]:has([data-testid="stFileUploaderDropzoneInstructions"]) > div {
    display: flex !important;
    flex-direction: column !important;
    align-items: center !important;
    width: 100% !important;
}
[data-testid="stFileUploaderDropzone"]:has([data-testid="stFileUploaderDropzoneInstructions"]) button {
    font-size: 0 !important;
    line-height: 0 !important;
    color: transparent !important;
    min-width: 110px !important;
    padding: 9px 22px !important;
    border-radius: 8px !important;
    border: 1px solid #E2E8F0 !important;
    background: #FFFFFF !important;
    transition: all 0.15s ease !important;
    cursor: pointer !important;
    position: relative !important;
    display: inline-flex !important;
    align-items: center !important;
    justify-content: center !important;
}
[data-testid="stFileUploaderDropzone"]:has([data-testid="stFileUploaderDropzoneInstructions"]) button:hover {
    border-color: #2563EB !important;
    background: #EFF6FF !important;
}
[data-testid="stFileUploaderDropzone"]:has([data-testid="stFileUploaderDropzoneInstructions"]) button::after {
    content: 'Chọn file' !important;
    font-size: 13px !important;
    font-weight: 600 !important;
    font-family: 'Inter', sans-serif !important;
    color: #2563EB !important;
    line-height: normal !important;
}
[data-testid="stFileUploaderDropzone"]:has([data-testid="stFileUploaderDropzoneInstructions"]) button > * {
    display: none !important;
}

/* Khi đã có file: ẩn hoàn toàn dropzone — không cần "Kéo thả" hay nút "Chọn file" nữa */
[data-testid="stFileUploaderDropzone"]:not(:has([data-testid="stFileUploaderDropzoneInstructions"])) {
    display: none !important;
}
[data-testid="stFileUploaderFile"] {
    position: absolute !important;
    bottom: 0 !important;
    left: 0 !important;
    right: 0 !important;
    padding: 8px 14px !important;
    overflow: hidden !important;
    display: flex !important;
    align-items: center !important;
    gap: 10px !important;
    flex-wrap: nowrap !important;
    background: #F8FAFC !important;
    border-radius: 0 !important;
    border: none !important;
    border-top: 1px solid #F1F5F9 !important;
    margin-bottom: 0 !important;
    width: auto !important;
    z-index: 1 !important;
}
[data-testid="stFileUploaderFile"] [data-testid="stFileUploaderFileIcon"] {
    flex-shrink: 0 !important;
    opacity: 0.5 !important;
}
[data-testid="stFileUploaderFile"] > div:first-child {
    overflow: hidden !important;
    text-overflow: ellipsis !important;
    white-space: nowrap !important;
    min-width: 0 !important;
    flex: 1 !important;
}
[data-testid="stFileUploaderFile"] small {
    overflow: hidden !important;
    text-overflow: ellipsis !important;
    white-space: nowrap !important;
    max-width: 160px !important;
    display: inline-block !important;
    font-family: 'JetBrains Mono', monospace !important;
    font-size: 12px !important;
    color: #334155 !important;
    max-width: 200px !important;
}
[data-testid="stFileUploaderFile"] button {
    font-size: inherit !important;
    min-width: unset !important;
    padding: 4px !important;
    border-radius: 50% !important;
    border: none !important;
    background: transparent !important;
    flex-shrink: 0 !important;
    opacity: 0.4 !important;
    transition: opacity 0.15s ease, background 0.15s ease !important;
}
[data-testid="stFileUploaderFile"] button:hover {
    opacity: 1 !important;
    background: rgba(220,38,38,0.06) !important;
}
[data-testid="stFileUploaderFile"] button::after { content: none !important; }

[data-testid="stFileUploaderDropzone"]:not(:has([data-testid="stFileUploaderDropzoneInstructions"])) > button,
[data-testid="stFileUploaderDropzone"]:not(:has([data-testid="stFileUploaderDropzoneInstructions"])) > span > button {
    font-size: 0 !important;
    line-height: 0 !important;
    color: transparent !important;
    padding: 6px 14px !important;
    border-radius: 8px !important;
    border: 1px solid #E2E8F0 !important;
    background: #FFFFFF !important;
    min-width: unset !important;
    cursor: pointer !important;
    transition: all 0.15s ease !important;
    display: inline-flex !important;
    align-items: center !important;
    justify-content: center !important;
    margin-top: 4px !important;
}
[data-testid="stFileUploaderDropzone"]:not(:has([data-testid="stFileUploaderDropzoneInstructions"])) > button:hover,
[data-testid="stFileUploaderDropzone"]:not(:has([data-testid="stFileUploaderDropzoneInstructions"])) > span > button:hover {
    border-color: #2563EB !important;
    background: #EFF6FF !important;
}
[data-testid="stFileUploaderDropzone"]:not(:has([data-testid="stFileUploaderDropzoneInstructions"])) > button::after,
[data-testid="stFileUploaderDropzone"]:not(:has([data-testid="stFileUploaderDropzoneInstructions"])) > span > button::after {
    content: 'Đổi file' !important;
    font-size: 12px !important;
    font-weight: 500 !important;
    font-family: 'Inter', sans-serif !important;
    color: #2563EB !important;
    line-height: normal !important;
}
[data-testid="stFileUploaderDropzone"]:not(:has([data-testid="stFileUploaderDropzoneInstructions"])) > button > *,
[data-testid="stFileUploaderDropzone"]:not(:has([data-testid="stFileUploaderDropzoneInstructions"])) > span > button > * {
    display: none !important;
}
[data-testid="stFileUploaderDropzone"]:not(:has([data-testid="stFileUploaderDropzoneInstructions"])) > span:has(button) {
    display: inline-flex !important; justify-content: center !important; width: auto !important;
}
[data-testid="stFileUploaderDropzone"]:not(:has([data-testid="stFileUploaderDropzoneInstructions"])) > span:nth-of-type(n+2) {
    display: none !important;
}

/* ════════════════════════════════════════════════════════════
   PLATFORM BADGE
════════════════════════════════════════════════════════════ */
.badge {
    display: inline-flex; align-items: center;
    font-size: 10.5px; font-weight: 600;
    padding: 3px 10px; border-radius: 6px;
    letter-spacing: 0.2px;
    font-family: 'Inter', sans-serif;
    vertical-align: middle;
    margin-left: 8px;
}
.badge-tiktok {
    background: #0F172A;
    color: #F1F5F9;
}
.badge-shopee {
    background: #FFF1F0;
    color: #DC2626;
    border: 1px solid rgba(220,38,38,0.15);
}

/* ════════════════════════════════════════════════════════════
   METRIC CARDS — Flat bento style
════════════════════════════════════════════════════════════ */
.os-metrics {
    display: grid;
    grid-template-columns: repeat(3, 1fr);
    gap: 10px;
    margin: 14px 0;
}
.os-metric {
    background: #F8FAFC;
    border: 1px solid #E2E8F0;
    border-radius: 12px;
    padding: 14px 16px;
    position: relative;
    transition: box-shadow 0.15s ease;
    display: flex;
    align-items: center;
    gap: 14px;
}
.os-metric:hover {
    box-shadow: 0 2px 12px rgba(0,0,0,0.06);
}
.os-metric-icon {
    width: 36px; height: 36px; flex-shrink: 0;
    background: #FFFFFF;
    border-radius: 9px;
    display: flex; align-items: center; justify-content: center;
    font-size: 16px;
    border: 1px solid #E2E8F0;
}
.os-metric-body { flex: 1; min-width: 0; }
.os-metric-label {
    font-size: 10px; font-weight: 600;
    color: #94A3B8; text-transform: uppercase;
    letter-spacing: 0.7px; margin-bottom: 2px;
}
.os-metric-value {
    font-size: 28px; font-weight: 800;
    color: #0F172A; letter-spacing: -1.5px;
    line-height: 1; margin-bottom: 2px;
    font-variant-numeric: tabular-nums;
}
.os-metric-note {
    font-size: 11px; color: #94A3B8; font-weight: 400;
}

/* ════════════════════════════════════════════════════════════
   DEDUP BLOCK — Flat alert card
════════════════════════════════════════════════════════════ */
.os-dedup {
    border: 1px solid #E2E8F0;
    border-radius: 12px;
    padding: 16px 20px;
    display: flex; align-items: center; gap: 16px;
    margin: 12px 0 20px;
}
.os-dedup.has-removed {
    background: #FFFBEB;
    border-color: #FDE68A;
    border-left: 3px solid #D97706;
}
.os-dedup.no-removed {
    background: #F0FDF4;
    border-color: #BBF7D0;
    border-left: 3px solid #16A34A;
}
.os-dedup-icon { font-size: 20px; flex-shrink: 0; line-height: 1; }
.os-dedup-body { flex: 1; }
.os-dedup-title {
    font-size: 13px; font-weight: 600; color: #0F172A;
    margin-bottom: 8px;
}
.os-dedup-pills { display: flex; gap: 5px; flex-wrap: wrap; }
.os-dedup-pill {
    font-size: 11px; font-weight: 500;
    font-family: 'Inter', sans-serif;
    padding: 3px 10px; border-radius: 100px;
    background: #F1F5F9; color: #64748B;
    border: 1px solid #E2E8F0;
}
.os-dedup-pill.removed {
    background: #FEF3C7; color: #92400E;
    border-color: #FDE68A;
}
.os-dedup-pill.kept {
    background: #DCFCE7; color: #166534;
    border-color: #BBF7D0;
}
.os-dedup-count {
    text-align: right; flex-shrink: 0;
    font-size: 32px; font-weight: 800;
    letter-spacing: -1.5px;
    font-variant-numeric: tabular-nums;
    line-height: 1;
}
.os-dedup-count.warn { color: #D97706; }
.os-dedup-count.ok   { color: #16A34A; }
.os-dedup-count-lbl {
    font-size: 11px; color: #94A3B8;
    margin-top: 3px; text-align: right; font-weight: 400;
}

/* ════════════════════════════════════════════════════════════
   SKU HEADER ROW
════════════════════════════════════════════════════════════ */
.os-sku-header {
    display: flex; align-items: center; justify-content: space-between;
    padding: 10px 0;
    border-bottom: 1px solid #F1F5F9;
    margin: 22px 0 6px;
}
.os-sku-left { display: flex; align-items: center; gap: 10px; }
.os-sku-code {
    font-family: 'JetBrains Mono', monospace;
    font-size: 11.5px; font-weight: 600;
    color: #2563EB;
    background: #EFF6FF;
    padding: 4px 12px; border-radius: 6px;
    letter-spacing: 0.2px;
    border: 1px solid #BFDBFE;
}
.os-sku-label { font-size: 12px; color: #94A3B8; font-weight: 400; }
.os-sku-total {
    font-size: 14px; font-weight: 700;
    color: #0F172A; letter-spacing: -0.3px;
    font-variant-numeric: tabular-nums;
}

/* ════════════════════════════════════════════════════════════
   DOWNLOAD BUTTON
════════════════════════════════════════════════════════════ */
/* ════════════════════════════════════════════════════════════
   DOWNLOAD BUTTON
════════════════════════════════════════════════════════════ */
.stDownloadButton > button {
    background: #2563EB !important;
    color: #FFFFFF !important;
    border: none !important;
    border-radius: 10px !important;
    padding: 11px 24px !important;
    font-weight: 600 !important;
    font-size: 13px !important;
    letter-spacing: -0.1px !important;
    font-family: 'Inter', sans-serif !important;
    box-shadow: 0 1px 3px rgba(37,99,235,0.2) !important;
    transition: background 0.15s ease, box-shadow 0.15s ease !important;
    width: 100% !important;
    height: 42px !important;
}
.stDownloadButton > button:hover {
    background: #1D4ED8 !important;
    box-shadow: 0 2px 8px rgba(37,99,235,0.25) !important;
}
.stDownloadButton > button:active {
    background: #1E40AF !important;
    box-shadow: 0 1px 3px rgba(37,99,235,0.15) !important;
}
.os-dl-row {
    display: flex;
    align-items: center;
    gap: 10px;
    padding: 10px 0;
    margin-top: 12px;
    flex-wrap: nowrap;
    min-width: 0;
}
.os-dl-download-btn-wrap {
    flex-shrink: 0;
}
.os-dl-download-btn-wrap .stDownloadButton > button {
    padding: 9px 18px !important;
    font-size: 13px !important;
    font-weight: 600 !important;
    border-radius: 9px !important;
    box-shadow: 0 1px 3px rgba(37,99,235,0.2) !important;
    white-space: nowrap !important;
    width: auto !important;
}
.os-dl-filename-pill {
    display: flex; align-items: center; gap: 6px;
    background: #F8FAFC;
    border: 1px solid #E2E8F0;
    border-radius: 8px;
    padding: 7px 12px;
    flex: 1; min-width: 0;
    overflow: hidden;
}
.os-dl-filename {
    font-size: 12px; font-weight: 500;
    color: #334155;
    font-family: 'Inter', sans-serif;
    white-space: nowrap;
    overflow: hidden;
    text-overflow: ellipsis;
    flex: 1;
    min-width: 0;
    letter-spacing: 0;
}
.os-dl-ext-badge {
    font-size: 10px; font-weight: 700;
    background: #DBEAFE; color: #1E40AF;
    padding: 2px 7px; border-radius: 5px;
    letter-spacing: 0.3px;
    flex-shrink: 0;
    font-family: 'Inter', sans-serif;
}
.os-dl-copy-btn {
    flex-shrink: 0;
    background: #FFFFFF;
    border: 1px solid #E2E8F0;
    border-radius: 8px;
    padding: 7px 13px;
    font-size: 12px; font-weight: 500;
    color: #475569;
    cursor: pointer;
    font-family: 'Inter', sans-serif;
    transition: all 0.15s ease;
    white-space: nowrap;
}
.os-dl-copy-btn:hover {
    border-color: #2563EB;
    color: #2563EB;
    background: #EFF6FF;
}
.os-dl-copy-btn.copied {
    background: #F0FDF4;
    border-color: #BBF7D0;
    color: #16A34A;
}
.os-dl-stats {
    font-size: 11.5px; color: #94A3B8;
    font-weight: 400; white-space: nowrap;
    flex-shrink: 0;
}

/* ════════════════════════════════════════════════════════════
   EMPTY STATE
════════════════════════════════════════════════════════════ */
.os-empty {
    text-align: center;
    padding: 56px 32px;
    border: 1px solid #F1F5F9;
    border-radius: 20px;
    background: #FAFAFA;
    margin: 12px 0 32px;
}
.os-empty-icon {
    font-size: 32px; margin-bottom: 16px; display: block;
    opacity: 0.5;
}
.os-empty-title {
    font-size: 17px; font-weight: 700; color: #0F172A;
    margin-bottom: 8px; letter-spacing: -0.3px;
}
.os-empty-sub {
    font-size: 13.5px; color: #94A3B8;
    max-width: 360px; margin: 0 auto 28px;
    line-height: 1.7; font-weight: 400;
}
.os-steps-row {
    display: flex; align-items: center;
    justify-content: center; gap: 6px;
    flex-wrap: wrap;
}
.os-step-item {
    display: flex; align-items: center; gap: 6px;
}
.os-step-sep {
    color: #CBD5E1; font-size: 12px;
}
.os-step-num {
    width: 22px; height: 22px;
    background: #F1F5F9;
    border-radius: 6px;
    display: inline-flex; align-items: center; justify-content: center;
    font-size: 11px; font-weight: 700; color: #64748B;
    font-family: 'JetBrains Mono', monospace;
    flex-shrink: 0;
}
.os-step-txt {
    font-size: 11.5px; color: #94A3B8; font-weight: 500;
}

/* ════════════════════════════════════════════════════════════
   UPLOAD NOTE
════════════════════════════════════════════════════════════ */
.os-note {
    font-size: 12px; color: #94A3B8;
    margin-top: 10px; line-height: 1.65;
    padding: 2px 0;
}
.os-note strong { color: #2563EB; font-weight: 600; }

/* ════════════════════════════════════════════════════════════
   DIVIDER
════════════════════════════════════════════════════════════ */
.os-divider {
    height: 1px;
    background: #F1F5F9;
    margin: 24px 0;
}

/* ════════════════════════════════════════════════════════════
   DATAFRAME
════════════════════════════════════════════════════════════ */
.stDataFrame {
    border: 1px solid #E2E8F0 !important;
    border-radius: 12px !important;
    overflow: hidden !important;
    box-shadow: none !important;
}
iframe { border-radius: 12px; }

/* ════════════════════════════════════════════════════════════
   RESPONSIVE
════════════════════════════════════════════════════════════ */
@media (max-width: 768px) {
    .os-topbar { flex-wrap: wrap; gap: 10px; }
    .os-topbar-status { margin-left: auto; }
    .os-metrics { grid-template-columns: 1fr; }
    .os-dedup { flex-direction: column; gap: 12px; }
    .os-dedup-count { text-align: left; }
    .os-dedup-count-lbl { text-align: left; }
    .os-steps-row { flex-direction: column; align-items: center; gap: 8px; }
    .os-step-sep { display: none; }
    .block-container { padding-left: 1.2rem !important; padding-right: 1.2rem !important; }
}
/* ════════════════════════════════════════════════════════════
   FILE UPLOADER — Refined layout
════════════════════════════════════════════════════════════ */
/* Wrap mỗi uploader trong card nhẹ */
[data-testid="stFileUploader"] {
    box-shadow: none !important;
}
/* Label căn trái, font bình thường */
[data-testid="stFileUploader"] label {
    padding: 12px 14px 10px !important;
    font-size: 13px !important;
    font-weight: 500 !important;
    color: #334155 !important;
    text-transform: none !important;
    text-align: left !important;
    letter-spacing: 0 !important;
}
/* Line phân cách giữa label và dropzone */
[data-testid="stFileUploader"] label::after {
    content: '';
    display: block;
    height: 1px;
    background: #F1F5F9;
    margin-top: 10px;
}
/* Dropzone inner padding thoáng hơn */
[data-testid="stFileUploaderDropzone"]:has([data-testid="stFileUploaderDropzoneInstructions"]) {
    padding: 18px 20px 22px !important;
}

/* ════════════════════════════════════════════════════════════
   GLOBAL FONT OVERRIDE — Inter everywhere
════════════════════════════════════════════════════════════ */
.stApp, .stApp * {
    font-family: 'Inter', system-ui, -apple-system, sans-serif !important;
}
/* Mono chỉ cho các element cụ thể */
.os-dl-filename, .os-sku-code, .os-dl-ext-badge,
[data-testid="stFileUploaderFile"] small {
    font-family: 'JetBrains Mono', monospace !important;
}
/* Streamlit default overrides */
[class*="stMarkdown"], [class*="stText"], p, span, div {
    font-family: 'Inter', system-ui, sans-serif;
}
/* Fix column gap giữa 2 ô upload */
[data-testid="stHorizontalBlock"] {
    gap: 12px !important;
}
/* Khi đã có file — wrapper đổi sang solid border */
[data-testid="stFileUploader"]:has([data-testid="stFileUploaderFile"]) {
    border-style: solid !important;
    border-color: #E2E8F0 !important;
    border-width: 1px !important;
    border-radius: 14px !important;
    overflow: hidden !important;
}
/* Nút đổi file — cạnh phải, nhỏ gọn */
[data-testid="stFileUploaderDropzone"]:not(:has([data-testid="stFileUploaderDropzoneInstructions"])) > button::after,
[data-testid="stFileUploaderDropzone"]:not(:has([data-testid="stFileUploaderDropzoneInstructions"])) > span > button::after {
    content: 'Đổi file' !important;
    font-size: 11px !important;
    font-weight: 500 !important;
    color: #64748B !important;
    line-height: normal !important;
}

/* ════ SPACING FIXES ════ */
/* Chips bắt buộc/tùy chọn — khoảng cách trên */
[data-testid="stVerticalBlock"] > [data-testid="stMarkdownContainer"]:has(div[style*="display:flex"]) {
    margin-top: 6px !important;
    margin-bottom: 4px !important;
}
/* Streamlit info/warning box — nhỏ gọn, đồng bộ font */
[data-testid="stAlert"] {
    margin-top: 10px !important;
    margin-bottom: 4px !important;
    border-radius: 8px !important;
    padding: 8px 12px !important;
}
[data-testid="stAlert"] p,
[data-testid="stAlert"] .stAlertContent,
[data-testid="stAlert"] div,
[data-testid="stAlert"] span {
    font-size: 11.5px !important;
    line-height: 1.5 !important;
    font-family: 'Inter', sans-serif !important;
    font-weight: 400 !important;
}

/* Upload columns — luôn cùng chiều cao (stretch) */
[data-testid="stHorizontalBlock"] {
    align-items: stretch !important;
}
[data-testid="stHorizontalBlock"] > [data-testid="stVerticalBlock"] {
    display: flex !important;
    flex-direction: column !important;
}
[data-testid="stHorizontalBlock"] > [data-testid="stVerticalBlock"] > [data-testid="stFileUploader"] {
    flex: 1 !important;
}

/* Upload wrapper: khi có file đổi sang solid border */
[data-testid="stFileUploader"]:has([data-testid="stFileUploaderFile"]) {
    border-style: solid !important;
    border-color: #E2E8F0 !important;
    border-width: 1px !important;
}

/* Section kết quả — tách rõ khỏi upload area */
.os-section {
    margin-top: 24px !important;
}
</style>
""", unsafe_allow_html=True)



# ─── HÀM TÁCH SKU ─────────────────────────────────────────────────────────────
def parse_sku_from_col_g(val):
    val = str(val).strip()
    return val.split('_')[0] if '_' in val else val

# ─── HÀM TÁCH MÀU & SIZE ──────────────────────────────────────────────────────
def parse_color_size_from_col_i(val):
    val = str(val).strip()
    if ',' in val:
        last_comma_index = val.rfind(',')
        color = val[:last_comma_index].strip()
        size = val[last_comma_index + 1:].strip() if last_comma_index < len(val) - 1 else "F"
        if ':' in size:
            size = size.split(':')[0].strip()
        color = ' '.join(color.replace(',', ' ').split())
    else:
        color = val
        size = "F"
    if ':' in size:
        size = size.split(':')[0].strip()
    return pd.Series([color, size])

# ─── HÀM HELPER FIELD CODE ────────────────────────────────────────────────────
def add_page_number_field(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    run._element.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._element.append(fldChar2)
    t = OxmlElement('w:t')
    run._element.append(t)
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar3)
    return run

def add_num_pages_field(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'NUMPAGES'
    run._element.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._element.append(fldChar2)
    t = OxmlElement('w:t')
    run._element.append(t)
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar3)
    return run

# ─── HÀM HELPER DOCX ──────────────────────────────────────────────────────────
def set_cell_shading(cell, color_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            element = OxmlElement(f'w:{edge}')
            element.set(qn('w:val'), val.get('val', 'single'))
            element.set(qn('w:sz'), val.get('sz', '4'))
            element.set(qn('w:color'), val.get('color', '000000'))
            element.set(qn('w:space'), '0')
            tcBorders.append(element)
    tcPr.append(tcBorders)

def set_row_height(row, height_cm=0.6):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)

def set_cell_margins(cell, top=30, bottom=30, left=60, right=60):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for edge, val in [('top', top), ('bottom', bottom), ('start', left), ('end', right)]:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)

def format_cell(cell, text, font_size=10, bold=False, align='left', font_name='Arial', color=None, indent=False):
    cell.text = ''
    para = cell.paragraphs[0]
    para_format = para.paragraph_format
    para_format.space_before = Pt(1)
    para_format.space_after = Pt(1)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    prefix = '    ' if indent else ''
    run = para.add_run(prefix + str(text))
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.name = font_name
    if color:
        run.font.color.rgb = RGBColor(*color)
    if align == 'center':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

# ─── HÀM XUẤT WORD ────────────────────────────────────────────────────────────
def export_to_word(detail_summary, total_orders, total_items, shop_name="TITIKID", platform="SHOPEE", shift="CHIỀU"):
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.clear()
    run_page = add_page_number_field(footer_para)
    run_page.font.size = Pt(9)
    run_page.font.name = 'Arial'
    run_sep = footer_para.add_run(' / ')
    run_sep.font.size = Pt(9)
    run_sep.font.name = 'Arial'
    run_total = add_num_pages_field(footer_para)
    run_total.font.size = Pt(9)
    run_total.font.name = 'Arial'

    current_date = datetime.now().strftime('%d.%m')
    title_text = f"{shop_name} - {platform} - {shift} - {current_date} - {total_orders} ĐƠN - {total_items} ÁO"
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title_text)
    title_run.font.size = Pt(13)
    title_run.font.bold = True
    title_run.font.name = 'Arial'
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    title_para.paragraph_format.space_after = Pt(6)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.autofit = False
    col_widths = [Inches(1.5), Inches(2.5), Inches(2.2), Inches(1.0)]

    hdr_cells = table.rows[0].cells
    headers = ['SKU sản phẩm', 'Màu', 'Size', 'ÁO']
    header_bg = 'C5D0DC'
    for i, header_text in enumerate(headers):
        format_cell(hdr_cells[i], header_text, font_size=10, bold=True, align='left' if i < 3 else 'right', font_name='Arial')
        set_cell_shading(hdr_cells[i], header_bg)
    set_row_height(table.rows[0], 0.9)

    top_total_cells = table.add_row().cells
    top_total_bg = 'DCE6F1'
    format_cell(top_total_cells[0], 'Tổng số', font_size=10, bold=True, font_name='Arial')
    format_cell(top_total_cells[1], '', font_size=10, font_name='Arial')
    format_cell(top_total_cells[2], '', font_size=10, font_name='Arial')
    format_cell(top_total_cells[3], '0', font_size=10, bold=True, align='right', font_name='Arial', color=(0, 0, 139))
    for cell in top_total_cells:
        set_cell_shading(cell, top_total_bg)
    set_row_height(table.rows[-1], 0.85)

    unique_skus = detail_summary['SKU'].unique()
    for sku in unique_skus:
        sku_data = detail_summary[detail_summary['SKU'] == sku].copy()
        total_sku = int(sku_data['SL'].sum())
        sku_data = sku_data.sort_values(by=['Phân loại', 'Size'])
        for _, data_row in sku_data.iterrows():
            row_cells = table.add_row().cells
            format_cell(row_cells[0], sku, font_size=10, bold=True, font_name='Arial')
            format_cell(row_cells[1], str(data_row['Phân loại']), font_size=10, bold=False, font_name='Arial')
            format_cell(row_cells[2], str(data_row['Size']), font_size=10, font_name='Arial')
            format_cell(row_cells[3], str(int(data_row['SL'])), font_size=10, align='right', font_name='Arial')
            for cell in row_cells:
                set_cell_shading(cell, 'FFFFFF')
            set_row_height(table.rows[-1], 0.85)

        total_row_cells = table.add_row().cells
        format_cell(total_row_cells[0], f'Tổng số {sku}', font_size=10, bold=True, font_name='Arial')
        format_cell(total_row_cells[1], '', font_size=10, font_name='Arial')
        format_cell(total_row_cells[2], '', font_size=10, font_name='Arial')
        format_cell(total_row_cells[3], str(total_sku), font_size=10, bold=True, align='right', font_name='Arial', color=(0, 0, 139))
        for cell in total_row_cells:
            set_cell_shading(cell, 'DCE6F1')
        set_row_height(table.rows[-1], 0.85)

    grand_total_cells = table.add_row().cells
    format_cell(grand_total_cells[0], 'Tổng cộng', font_size=10, bold=True, font_name='Arial')
    format_cell(grand_total_cells[1], '', font_size=10, font_name='Arial')
    format_cell(grand_total_cells[2], '', font_size=10, font_name='Arial')
    format_cell(grand_total_cells[3], str(total_items), font_size=10, bold=True, align='right', font_name='Arial', color=(0, 0, 139))
    for cell in grand_total_cells:
        set_cell_shading(cell, 'DCE6F1')
    set_row_height(table.rows[-1], 0.85)

    for row in table.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = w

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    for row_idx, row in enumerate(table.rows):
        if row_idx == 0:
            trPr = row._tr.get_or_add_trPr()
            tblHeader = OxmlElement('w:tblHeader')
            trPr.append(tblHeader)

    target = io.BytesIO()
    doc.save(target)
    return target.getvalue()

# ══════════════════════════════════════════════════════════════════════════════
#  TIÊU ĐỀ
# ══════════════════════════════════════════════════════════════════════════════
st.markdown('''
<div class="os-topbar">
    <div class="os-topbar-brand">
        <div class="os-topbar-logo">📦</div>
        <div>
            <div class="os-topbar-name">Order Studio</div>
            <div class="os-topbar-sub">TikTok & Shopee · Tổng hợp · Lọc trùng · Xuất Word</div>
        </div>
    </div>
    <div class="os-topbar-status">
        <span class="os-status-dot"></span> Sẵn sàng
    </div>
</div>
''', unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
#  01 — CÀI ĐẶT
# ══════════════════════════════════════════════════════════════════════════════
st.markdown('''
<div class="os-section">
    <span class="os-section-title">Cài đặt</span>
    <div class="os-section-line"></div>
</div>
''', unsafe_allow_html=True)


current_hour = datetime.now().hour
default_shift_index = 0 if current_hour < 12 else 1

col_s1, col_s2, col_s3 = st.columns(3)
with col_s1:
    shop_name = st.selectbox("Tên shop", ["TITIKID", "GIMME"], key="shop_name")
with col_s2:
    platform = st.selectbox("Sàn bán hàng", ["TIKTOK", "SHOPEE"], key="platform")
with col_s3:
    shift = st.selectbox("Ca làm việc", ["SÁNG", "CHIỀU"], index=default_shift_index, key="shift")


# ══════════════════════════════════════════════════════════════════════════════
#  02 — UPLOAD FILE
# ══════════════════════════════════════════════════════════════════════════════
platform_badge = (
    '<span class="badge badge-tiktok">TikTok</span>'
    if platform == "TIKTOK"
    else '<span class="badge badge-shopee">Shopee</span>'
)

st.markdown(f'''
<div class="os-section">
    <span class="os-section-title">Tải file đơn hàng {platform_badge}</span>
    <div class="os-section-line"></div>
</div>
''', unsafe_allow_html=True)

col_f1, col_f2 = st.columns(2)
with col_f1:
    uploaded_file = st.file_uploader(
        f"File ca hiện tại ({platform})",
        type=["csv", "xlsx"],
        key="file_current",
        help="File đơn hàng ca này cần tổng hợp"
    )
with col_f2:
    prev_file = st.file_uploader(
        "File ca trước (tùy chọn)",
        type=["csv", "xlsx"],
        key="file_prev",
        help="Tải lên để loại bỏ đơn đã soạn ở ca trước"
    )

st.markdown('''
<div style="display:flex; gap:10px; margin-top:10px; align-items:center; flex-wrap:wrap;">
    <div style="display:inline-flex; align-items:center; gap:6px; background:#EFF6FF; border:1px solid #BFDBFE; border-radius:8px; padding:7px 13px;">
        <svg width="13" height="13" viewBox="0 0 16 16" fill="none"><circle cx="8" cy="8" r="7" stroke="#2563EB" stroke-width="1.5"/><path d="M8 5v4M8 11v.5" stroke="#2563EB" stroke-width="1.5" stroke-linecap="round"/></svg>
        <span style="font-size:12.5px; font-weight:500; color:#1E40AF; font-family:Inter,sans-serif;">File ca hiện tại — <strong style="font-weight:700;">Bắt buộc</strong></span>
    </div>
    <div style="display:inline-flex; align-items:center; gap:6px; background:#F8FAFC; border:1px solid #E2E8F0; border-radius:8px; padding:7px 13px;">
        <svg width="13" height="13" viewBox="0 0 16 16" fill="none"><circle cx="8" cy="8" r="7" stroke="#94A3B8" stroke-width="1.5"/><path d="M8 5v4M8 11v.5" stroke="#94A3B8" stroke-width="1.5" stroke-linecap="round"/></svg>
        <span style="font-size:12.5px; font-weight:500; color:#64748B; font-family:Inter,sans-serif;">File ca trước — <span style="color:#94A3B8;">Tùy chọn · dùng để lọc đơn trùng</span></span>
    </div>
</div>
''', unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
#  03 — KẾT QUẢ
# ══════════════════════════════════════════════════════════════════════════════
if not uploaded_file:
    st.markdown('''
    <div class="os-section" style="margin-top:28px;">
        <span class="os-section-title">Kết quả</span>
        <div class="os-section-line"></div>
    </div>
    <div class="os-empty">
        <span class="os-empty-icon">📦</span>
        <div class="os-empty-title">Chưa có dữ liệu</div>
        <div class="os-empty-sub">Tải lên file đơn hàng phía trên để bắt đầu tổng hợp, lọc trùng và xuất file Word soạn hàng.</div>
        <div class="os-steps-row">
            <div class="os-step-item">
                <div class="os-step-num">1</div>
                <div class="os-step-txt">Chọn cửa hàng</div>
            </div>
            <span class="os-step-sep">→</span>
            <div class="os-step-item">
                <div class="os-step-num">2</div>
                <div class="os-step-txt">Tải file lên</div>
            </div>
            <span class="os-step-sep">→</span>
            <div class="os-step-item">
                <div class="os-step-num">3</div>
                <div class="os-step-txt">Lọc trùng</div>
            </div>
            <span class="os-step-sep">→</span>
            <div class="os-step-item">
                <div class="os-step-num">4</div>
                <div class="os-step-txt">Tải Word</div>
            </div>
        </div>
    </div>
    ''', unsafe_allow_html=True)

else:
    try:
        # ── Đọc file hiện tại ────────────────────────────────────────────────
        if uploaded_file.name.endswith('.csv'):
            df = pd.read_csv(uploaded_file, low_memory=False, dtype=str)
        else:
            df = pd.read_excel(uploaded_file, engine='calamine', dtype=str)
        df = df.dropna(how='all').reset_index(drop=True)

        # ── Mapping cột theo sàn ─────────────────────────────────────────────
        if platform == "TIKTOK":
            sku_col_index, variation_col_index, qty_col_index = 6, 8, 9
        else:
            # Shopee: cột khác nhau theo shop
            if shop_name == "GIMME":
                # GIMME: P(15)=SKU, V(21)=Phân loại, AB(27)=Số lượng
                sku_col_index, variation_col_index, qty_col_index = 15, 21, 27
            else:
                # TITIKID: S(18)=SKU, T(19)=Phân loại, Z(25)=Số lượng
                sku_col_index, variation_col_index, qty_col_index = 18, 19, 25

        max_col_needed = max(sku_col_index, variation_col_index, qty_col_index)
        if len(df.columns) <= max_col_needed:
            st.error(f"❌ File không đủ cột cho sàn {platform}. Cần ít nhất {max_col_needed + 1} cột, file chỉ có {len(df.columns)} cột.")
            st.stop()

        col_sku       = df.columns[sku_col_index]
        col_variation = df.columns[variation_col_index]
        col_qty       = df.columns[qty_col_index]

        # Shopee GIMME dùng cột A để đếm/lọc đơn; Shopee TITIKID dùng cột G (mã vận đơn).
        if platform == "SHOPEE" and shop_name != "GIMME":
            if len(df.columns) <= 6:
                st.error("❌ File Shopee không đủ cột. Cần ít nhất cột G (mã vận đơn).")
                st.stop()
            id_col = df.columns[6]  # Cột G = mã vận đơn
            # Bỏ những đơn không có mã vận đơn
            no_tracking_count = df[id_col].isna().sum() + (df[id_col].astype(str).str.strip() == '').sum()
            df = df[df[id_col].astype(str).str.strip().ne('') & df[id_col].notna()].reset_index(drop=True)
            if no_tracking_count > 0:
                st.info(f"ℹ️ Đã bỏ **{no_tracking_count}** đơn không có mã vận đơn (cột G trống).")
        else:
            id_col = df.columns[0]  # TikTok: vẫn dùng cột A

        total_raw     = df[id_col].nunique()
        removed_count = 0
        prev_total    = 0

        # ── Nhãn mục 03 ─────────────────────────────────────────────────
        st.markdown('''
        <div class="os-section" style="margin-top:28px;">
            <span class="os-section-title">Kết quả</span>
            <div class="os-section-line"></div>
            <span class="os-section-tag">xuất dữ liệu</span>
        </div>
        ''', unsafe_allow_html=True)

        # ── Lọc trùng với file buổi trước ───────────────────────────────────
        if prev_file is not None:
            try:
                if prev_file.name.endswith('.csv'):
                    df_prev = pd.read_csv(prev_file, low_memory=False, dtype=str)
                else:
                    df_prev = pd.read_excel(prev_file, engine='calamine', dtype=str)
                df_prev = df_prev.dropna(how='all').reset_index(drop=True)

                # Shopee GIMME dùng cột A cho file ca trước; Shopee TITIKID dùng cột G.
                if platform == "SHOPEE" and shop_name != "GIMME" and len(df_prev.columns) > 6:
                    prev_id_col = df_prev.columns[6]
                    # Bỏ dòng không có mã vận đơn trong file ca trước
                    df_prev = df_prev[df_prev[prev_id_col].astype(str).str.strip().ne('') & df_prev[prev_id_col].notna()].reset_index(drop=True)
                else:
                    prev_id_col = df_prev.columns[0]

                prev_ids     = set(df_prev[prev_id_col].dropna().astype(str).str.strip())
                current_ids  = set(df[id_col].dropna().astype(str).str.strip())
                duplicated   = current_ids & prev_ids
                removed_count = len(duplicated)
                prev_total   = len(prev_ids)

                df = df[~df[id_col].astype(str).str.strip().isin(duplicated)].reset_index(drop=True)
                kept_count = df[id_col].nunique()

                if removed_count > 0:
                    st.markdown(f'''
                    <div class="os-dedup has-removed">
                        <div class="os-dedup-icon">⇄</div>
                        <div class="os-dedup-body">
                            <div class="os-dedup-title">Đã lọc trùng với file ca trước</div>
                            <div class="os-dedup-pills">
                                <span class="os-dedup-pill">Ca này: {total_raw} đơn</span>
                                <span class="os-dedup-pill">Ca trước: {prev_total} đơn</span>
                                <span class="os-dedup-pill removed">Loại: {removed_count} trùng</span>
                                <span class="os-dedup-pill kept">Giữ: {kept_count} mới</span>
                            </div>
                        </div>
                        <div>
                            <div class="os-dedup-count warn">−{removed_count}</div>
                            <div class="os-dedup-count-lbl">đơn trùng</div>
                        </div>
                    </div>
                    ''', unsafe_allow_html=True)
                else:
                    st.markdown(f'''
                    <div class="os-dedup no-removed">
                        <div class="os-dedup-icon">✓</div>
                        <div class="os-dedup-body">
                            <div class="os-dedup-title">Không có đơn trùng</div>
                            <div class="os-dedup-pills">
                                <span class="os-dedup-pill">Ca này: {total_raw} đơn</span>
                                <span class="os-dedup-pill">Ca trước: {prev_total} đơn</span>
                                <span class="os-dedup-pill kept">Tất cả đều là đơn mới</span>
                            </div>
                        </div>
                        <div>
                            <div class="os-dedup-count ok">0</div>
                            <div class="os-dedup-count-lbl">đơn trùng</div>
                        </div>
                    </div>
                    ''', unsafe_allow_html=True)
            except Exception as e_prev:
                st.warning(f"⚠️ Không đọc được file buổi trước: {e_prev}")

        # ── Xử lý data ───────────────────────────────────────────────────────
        df['SKU_ID'] = df[col_sku].apply(parse_sku_from_col_g)
        df[['PL', 'SZ']] = df[col_variation].apply(parse_color_size_from_col_i)
        df['SL'] = pd.to_numeric(df[col_qty], errors='coerce').fillna(0).astype(int)

        total_items  = int(df['SL'].sum())
        total_orders = df[id_col].nunique()

        detail_summary = df.groupby(['SKU_ID', 'PL', 'SZ'])['SL'].sum().reset_index()
        detail_summary.columns = ['SKU', 'Phân loại', 'Size', 'SL']
        unique_skus_count = detail_summary['SKU'].nunique()

        # ── Metric Cards ─────────────────────────────────────────────────────
        dedup_note = f"sau khi lọc {removed_count} trùng" if removed_count > 0 else "đơn hàng duy nhất"
        st.markdown(f'''
        <div class="os-metrics">
            <div class="os-metric">
                <div class="os-metric-icon">📦</div>
                <div class="os-metric-body">
                    <div class="os-metric-label">Tổng đơn hàng</div>
                    <div class="os-metric-value">{total_orders}</div>
                    <div class="os-metric-note">{dedup_note}</div>
                </div>
            </div>
            <div class="os-metric">
                <div class="os-metric-icon">👕</div>
                <div class="os-metric-body">
                    <div class="os-metric-label">Tổng sản phẩm</div>
                    <div class="os-metric-value">{total_items}</div>
                    <div class="os-metric-note">tổng số lượng áo</div>
                </div>
            </div>
            <div class="os-metric">
                <div class="os-metric-icon">🏷️</div>
                <div class="os-metric-body">
                    <div class="os-metric-label">Loại SKU</div>
                    <div class="os-metric-value">{unique_skus_count}</div>
                    <div class="os-metric-note">mã sản phẩm khác nhau</div>
                </div>
            </div>
        </div>
        ''', unsafe_allow_html=True)

        # ── Download Word ─────────────────────────────────────────────────────
        current_date_file = datetime.now().strftime('%d.%m')
        word_filename = f"{shop_name}_{platform}_{shift}_{current_date_file}_{total_orders} ĐƠN_{total_items} ÁO.docx"
        word_filename_display = f"{shop_name} {platform} {shift} {current_date_file} {total_orders} ĐƠN {total_items} ÁO"
        word_data = export_to_word(detail_summary, total_orders, total_items, shop_name, platform, shift)

        # ── Download row ──
        st.markdown(f'''
        <div class="os-dl-row">
            <div class="os-dl-filename-pill">
                <span class="os-dl-filename" id="os-filename-text">{word_filename_display}</span>
                <span class="os-dl-ext-badge">.DOCX</span>
            </div>
            <button class="os-dl-copy-btn" id="os-copy-btn" onclick="
                var txt = document.getElementById(\'os-filename-text\').innerText;
                navigator.clipboard.writeText(txt).then(function(){{
                    var btn = document.getElementById(\'os-copy-btn\');
                    btn.innerHTML = \'✓ Đã copy\';
                    btn.classList.add(\'copied\');
                    setTimeout(function(){{ btn.innerHTML = \'⎘ Copy tên\'; btn.classList.remove(\'copied\'); }}, 2000);
                }});
            ">⎘ Copy tên</button>
            <span class="os-dl-stats">{total_orders} đơn &middot; {total_items} sp &middot; {unique_skus_count} SKU</span>
        </div>
        ''', unsafe_allow_html=True)
        st.download_button(
            f"⬇️  Tải xuống phiếu soạn hàng  —  {total_orders} đơn · {total_items} áo",
            word_data,
            word_filename,
            use_container_width=True
        )

        # ── Chi tiết theo SKU ─────────────────────────────────────────────────
        st.markdown('<div class="os-divider"></div>', unsafe_allow_html=True)

        unique_skus = detail_summary['SKU'].unique()
        for sku in unique_skus:
            sku_data  = detail_summary[detail_summary['SKU'] == sku].sort_values(by='Size')
            total_sku = int(sku_data['SL'].sum())

            st.markdown(f'''
            <div class="os-sku-header">
                <div class="os-sku-left">
                    <span class="os-sku-code">{sku}</span>
                    <span class="os-sku-label">sản phẩm</span>
                </div>
                <span class="os-sku-total">{total_sku} cái</span>
            </div>
            ''', unsafe_allow_html=True)

            st.dataframe(
                sku_data[['Phân loại', 'Size', 'SL']],
                use_container_width=True,
                hide_index=True,
                column_config={
                    "Phân loại": st.column_config.TextColumn("Phân loại / Màu", width="large"),
                    "Size": st.column_config.TextColumn("Size", width="medium"),
                    "SL": st.column_config.NumberColumn("SL", width="small", format="%d")
                }
            )

    except Exception as e:
        st.error(f"❌ Lỗi xử lý file: {e}")
